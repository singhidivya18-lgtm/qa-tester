"""DOCX report generator tool.

Converts the QA report text into a professionally formatted .docx file.
Supports embedding screenshots from the testing run.
"""

import os
import json
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .screenshot_tool import SCREENSHOT_DIR, get_screenshot_dir


def generate_docx_report(report_text, all_screen_results_json="[]", site_url="", screenshots_dir=None, output_dir=None):
    """Generate a formatted .docx file from report text with embedded screenshots.

    Args:
        report_text: The full markdown/text report from the report agent.
        all_screen_results_json: JSON string of all screen results with verdicts.
        site_url: The tested site URL.
        screenshots_dir: Directory containing screenshot PNGs (defaults to SCREENSHOT_DIR).
        output_dir: Directory to save the .docx file (defaults to pipeline root).

    Returns:
        str: Path to the generated .docx file.
    """
    if screenshots_dir is None:
        screenshots_dir = SCREENSHOT_DIR
    if output_dir is None:
        output_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    try:
        screen_results = json.loads(all_screen_results_json) if all_screen_results_json else []
    except (json.JSONDecodeError, TypeError):
        screen_results = []

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    title = doc.add_heading("React QA Test Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    timestamp_paragraph = doc.add_paragraph()
    timestamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = timestamp_paragraph.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    if site_url:
        url_paragraph = doc.add_paragraph()
        url_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        url_run = url_paragraph.add_run(f"Site: {site_url}")
        url_run.font.size = Pt(10)
        url_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    lines = report_text.split("\n")
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                if code_lines:
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                code_lines = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 50)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            text = _clean_markdown(text)
            doc.add_paragraph(text, style="List Bullet")
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in (".", ")"):
            text = stripped[2:].strip()
            text = _clean_markdown(text)
            doc.add_paragraph(text, style="List Number")
        elif "|" in stripped and stripped.startswith("|"):
            text = _clean_markdown(stripped)
            text = text.replace("|", "  •  ")
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(9)
        else:
            text = _clean_markdown(stripped)
            p = doc.add_paragraph(text)
            _apply_inline_formatting(p, stripped)

    doc.add_page_break()
    _add_results_section(doc, screen_results)
    doc.add_heading("Screenshots", level=1)
    _add_screenshots_section(doc, screen_results, screenshots_dir)

    os.makedirs(output_dir, exist_ok=True)
    timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"qa_report_{timestamp_str}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return filepath


def _add_results_section(doc, screen_results):
    """Per-screen verdict table — always built from the results JSON (deterministic)."""
    if not screen_results:
        doc.add_paragraph("No check results recorded for any screen.")
        return
    for r in screen_results:
        if not isinstance(r, dict):
            continue
        sid = r.get("screen_id", "?")
        overall = r.get("overall", "?")
        route = r.get("route", "")
        doc.add_heading(f"Screen: {sid} — Overall: {overall}", level=2)
        if route:
            doc.add_paragraph(f"Route: {route}")
        verdicts = r.get("verdicts", [])
        if not verdicts:
            doc.add_paragraph("No verdicts recorded for this screen.")
            continue
        try:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
        except Exception:
            table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        for i, h in enumerate(("Check", "Verdict", "Detail", "Screenshot")):
            hdr[i].text = h
        for v in verdicts:
            if not isinstance(v, dict):
                continue
            verdict = v.get("verdict", "?")
            label = "BROKEN" if verdict == "FAIL" else verdict
            row = table.add_row().cells
            row[0].text = str(v.get("check_id", "?"))
            row[1].text = label
            row[2].text = (v.get("reason") or "")[:400]
            row[3].text = os.path.basename(v.get("screenshot_path", ""))
        doc.add_paragraph()


def _add_screenshots_section(doc, screen_results, screenshots_dir):
    """Embed screenshots for each screen with a verdict caption above each."""
    if not os.path.isdir(screenshots_dir):
        doc.add_paragraph("No screenshots directory found — screenshots were not saved during testing.")
        return

    screenshot_files = sorted(
        [f for f in os.listdir(screenshots_dir) if f.endswith(".png")],
        key=lambda f: os.path.getmtime(os.path.join(screenshots_dir, f)),
    )

    if not screenshot_files:
        doc.add_paragraph("No screenshots were captured during testing.")
        return

    # Lookup: screenshot basename -> (screen_id, verdict) so each image gets a caption
    by_path = {}
    for r in screen_results:
        if not isinstance(r, dict):
            continue
        for v in r.get("verdicts", []):
            if isinstance(v, dict) and v.get("screenshot_path"):
                by_path[os.path.basename(v["screenshot_path"])] = (r.get("screen_id", "?"), v)

    # Group screenshots per screen (verdict screen_id when matched, else filename prefix)
    groups = {}
    for fname in screenshot_files:
        info = by_path.get(fname)
        if info:
            sid = info[0]
        else:
            parts = fname.replace(".png", "").split("_")
            sid = parts[0] + "_" + parts[1] if len(parts) >= 2 else "other"
        groups.setdefault(sid, []).append(fname)

    TITLES = {"login": "Login", "register": "Registration"}
    for sid, files in sorted(groups.items()):
        doc.add_heading(TITLES.get(sid, f"Screen: {sid}"), level=2)

        for fname in files[:4]:
            fpath = os.path.join(screenshots_dir, fname)
            info = by_path.get(fname)
            if info:
                _, v = info
                label = "BROKEN" if v.get("verdict") == "FAIL" else v.get("verdict", "?")
                reason = (v.get("reason") or "").strip()[:300]
                caption = f"{v.get('check_id', '?')} — {label} — {reason}"
            else:
                caption = f"{fname} — no verdict recorded for this screenshot"
            p = doc.add_paragraph()
            p_run = p.add_run(caption)
            p_run.font.size = Pt(9)
            p_run.bold = True
            try:
                doc.add_picture(fpath, width=Inches(5.5))
            except Exception:
                doc.add_paragraph(f"[Could not embed: {fname}]")

        if len(files) > 4:
            doc.add_paragraph(f"(+ {len(files) - 4} more screenshots for this screen)")


def _clean_markdown(text):
    import re
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _apply_inline_formatting(paragraph, text):
    import re
    paragraph.clear()
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


